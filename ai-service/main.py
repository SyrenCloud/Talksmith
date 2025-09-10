from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import List, Optional
import os
import tempfile
from dotenv import load_dotenv
import asyncio
import aiofiles
import logging
import time
import json
from functools import wraps
import shutil as shutil_module

# LangChain imports
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_chroma import Chroma
import chromadb
from chromadb.config import Settings

# Document processing
import pdfplumber
import docx
import io
from uuid import uuid4
import shutil

# Cross-encoder for reranking
from sentence_transformers import CrossEncoder

# Load environment variables
load_dotenv()

# Configure logging for production
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

app = FastAPI(
    title="Talksmith AI Service", 
    version="1.0.0",
    description="Production-grade RAG system with multi-document support"
)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize LLM with GPT-4 Turbo (best available model)
def initialize_llm():
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise ValueError("OPENAI_API_KEY not found in environment variables")
    
    return ChatOpenAI(
        model_name=os.getenv("MODEL_NAME", "gpt-4"),
        temperature=float(os.getenv("TEMPERATURE", "1.0")),
        openai_api_key=api_key
    )

def initialize_embeddings():
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise ValueError("OPENAI_API_KEY not found in environment variables")
    
    return OpenAIEmbeddings(
        openai_api_key=api_key,
        model="text-embedding-3-large"
    )

# Initialize models
llm = initialize_llm()
embeddings = initialize_embeddings()

# Initialize cross-encoder for reranking
cross_encoder = CrossEncoder('cross-encoder/ms-marco-MiniLM-L12-v2')
logger.info("Cross-encoder initialized for document reranking")

# Pydantic models
class ProcessDocumentsRequest(BaseModel):
    file_paths: List[str]

class ChatRequest(BaseModel):
    message: str
    documents: List[str]
    conversationHistory: Optional[List[dict]] = []

class ChatResponse(BaseModel):
    response: str
    sources: Optional[List[str]] = None

class ProcessResponse(BaseModel):
    document_ids: List[str]
    message: str

class ProcessProgressEvent(BaseModel):
    type: str  # 'progress', 'stage', 'complete', 'error'
    message: str
    progress: int  # 0-100
    current_file: Optional[str] = None
    total_files: Optional[int] = None
    processed_files: Optional[int] = None

# Initialize ChromaDB client with persistence
PERSIST_DIRECTORY = os.path.join(os.path.dirname(__file__), "chroma_db")

# Create persist directory if it doesn't exist
os.makedirs(PERSIST_DIRECTORY, exist_ok=True)

# Initialize ChromaDB client
chroma_client = chromadb.PersistentClient(
    path=PERSIST_DIRECTORY,
    settings=Settings(
        anonymized_telemetry=False,
        allow_reset=True
    )
)

# Production-ready ChromaDB configuration
class ChromaDBConfig:
    """Configuration for ChromaDB collections"""
    MAIN_COLLECTION = "talksmith_documents"  # Single collection for all documents
    MAX_BATCH_SIZE = 100  # ChromaDB batch size limit
    SEARCH_TYPE = "similarity"  # or "mmr" for diversity
    SEARCH_KWARGS = {"k": 10}  # Increased k since we're searching across all documents
    DISTANCE_METRIC = "cosine"  # ChromaDB distance metric
    
    # Reranking configuration
    RERANK_MULTIPLIER = 2  # Retrieve this many times more candidates for reranking
    RERANK_TOP_K = 5  # Final number of chunks after reranking
    MIN_RERANK_SCORE = 0.3  # Minimum cross-encoder score to include a chunk

# Document processing functions
def load_document(file_path: str) -> List[Document]:
    """Load document based on file type"""
    documents = []
    
    try:
        if file_path.lower().endswith('.pdf'):
            with pdfplumber.open(file_path) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:  # Only add non-empty pages
                        text += page_text + "\n"
                documents.append(Document(page_content=text, metadata={"source": file_path}))
                
        elif file_path.lower().endswith('.docx'):
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            documents.append(Document(page_content=text, metadata={"source": file_path}))
            
        elif file_path.lower().endswith(('.txt', '.md')):
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
                documents.append(Document(page_content=text, metadata={"source": file_path}))
                
        else:
            raise ValueError(f"Unsupported file type: {file_path}")
            
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error processing file {file_path}: {str(e)}")
    
    return documents

def get_main_vectorstore() -> Chroma:
    """Get or create the main ChromaDB vector store"""
    try:
        # Check if collection exists
        collections = chroma_client.list_collections()
        collection_exists = any(col.name == ChromaDBConfig.MAIN_COLLECTION for col in collections)
        
        if not collection_exists:
            logger.info(f"Creating new collection: {ChromaDBConfig.MAIN_COLLECTION}")
        
        return Chroma(
            client=chroma_client,
            collection_name=ChromaDBConfig.MAIN_COLLECTION,
            embedding_function=embeddings,
            persist_directory=PERSIST_DIRECTORY,
            collection_metadata={"hnsw:space": ChromaDBConfig.DISTANCE_METRIC}
        )
    except Exception as e:
        logger.error(f"Error accessing main collection: {str(e)}")
        raise

def process_documents(documents: List[Document], document_id: str) -> int:
    """Split documents into chunks and add to main ChromaDB collection"""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len,
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    
    chunks = text_splitter.split_documents(documents)
    
    # Add document_id to metadata for each chunk
    for chunk in chunks:
        chunk.metadata["document_id"] = document_id
    
    # Get main vector store
    vectorstore = get_main_vectorstore()
    
    # Add documents in batches to avoid ChromaDB limits
    batch_size = ChromaDBConfig.MAX_BATCH_SIZE
    for i in range(0, len(chunks), batch_size):
        batch = chunks[i:i + batch_size]
        vectorstore.add_documents(documents=batch)
        logger.info(f"Added batch {i//batch_size + 1} of {len(chunks)//batch_size + 1} for document {document_id}")
    
    logger.info(f"Added {len(chunks)} chunks from document {document_id} to main collection")
    
    return len(chunks)

# Global storage for document metadata
document_metadata = {}  # Track document processing stats

# Performance monitoring decorator
def monitor_performance(func):
    @wraps(func)
    async def wrapper(*args, **kwargs):
        start_time = time.time()
        try:
            result = await func(*args, **kwargs)
            duration = time.time() - start_time
            logger.info(f"{func.__name__} completed in {duration:.2f}s")
            return result
        except Exception as e:
            duration = time.time() - start_time
            logger.error(f"{func.__name__} failed after {duration:.2f}s: {str(e)}")
            raise
    return wrapper

# Routes
@app.get("/")
async def root():
    return {"message": "Talksmith AI Service is running", "version": "1.0.0"}

@app.get("/health")
async def health_check():
    """Health check endpoint with ChromaDB status"""
    try:
        # Check ChromaDB connection
        collections = chroma_client.list_collections()
        collection_count = len(collections)
        
        # Get chunk count from main collection if it exists
        total_chunks = 0
        try:
            collection = chroma_client.get_collection(name=ChromaDBConfig.MAIN_COLLECTION)
            total_chunks = collection.count()
        except:
            pass
        
        return {
            "status": "healthy",
            "llm_model": os.getenv("MODEL_NAME", "gpt-4"),
            "vector_store": "ChromaDB",
            "persist_directory": PERSIST_DIRECTORY,
            "main_collection": ChromaDBConfig.MAIN_COLLECTION,
            "collection_count": collection_count,
            "active_documents": len(document_metadata),
            "total_chunks": total_chunks
        }
    except Exception as e:
        logger.error(f"Health check failed: {str(e)}")
        return {
            "status": "unhealthy",
            "error": str(e)
        }

@app.post("/process-stream")
async def process_documents_stream(request: ProcessDocumentsRequest):
    """Process documents with real-time progress updates via SSE"""
    
    async def generate_progress():
        try:
            document_ids = []
            total_files = len(request.file_paths)
            processed_files = 0
            
            # Send initial progress
            yield f"data: {json.dumps({'type': 'stage', 'message': 'Starting document processing...', 'progress': 0, 'total_files': total_files, 'processed_files': 0})}\n\n"
            
            # Process each document
            for i, file_path in enumerate(request.file_paths):
                file_name = os.path.basename(file_path)
                
                # Send current file progress
                yield f"data: {json.dumps({'type': 'progress', 'message': f'Processing {file_name}...', 'progress': int((i / total_files) * 80), 'current_file': file_name, 'total_files': total_files, 'processed_files': processed_files})}\n\n"
                
                if not os.path.exists(file_path):
                    logger.warning(f"File not found: {file_path}")
                    yield f"data: {json.dumps({'type': 'progress', 'message': f'File not found: {file_name}', 'progress': int(((i + 1) / total_files) * 80), 'current_file': file_name, 'total_files': total_files, 'processed_files': processed_files})}\n\n"
                    document_ids.append(None)
                    continue
                
                try:
                    # Load document
                    yield f"data: {json.dumps({'type': 'stage', 'message': f'Loading {file_name}...', 'progress': int((i / total_files) * 80) + 5, 'current_file': file_name, 'total_files': total_files, 'processed_files': processed_files})}\n\n"
                    docs = load_document(file_path)
                    
                    if docs:
                        # Create embeddings
                        yield f"data: {json.dumps({'type': 'stage', 'message': f'Creating embeddings for {file_name}...', 'progress': int((i / total_files) * 80) + 10, 'current_file': file_name, 'total_files': total_files, 'processed_files': processed_files})}\n\n"
                        
                        # Generate unique ID
                        doc_id = str(uuid4())
                        
                        # Process document into main ChromaDB collection
                        chunk_count = process_documents(docs, doc_id)
                        
                        # Store metadata
                        document_metadata[doc_id] = {
                            "filename": file_name,
                            "processed_at": time.time(),
                            "chunk_count": chunk_count,
                            "file_size": os.path.getsize(file_path)
                        }
                        document_ids.append(doc_id)
                        processed_files += 1
                        
                        # Move processed file to processed-documents folder
                        try:
                            # Create processed-documents directory if it doesn't exist
                            uploads_dir = os.path.dirname(file_path)
                            processed_dir = os.path.join(uploads_dir, "processed-documents")
                            os.makedirs(processed_dir, exist_ok=True)
                            
                            # Move the file
                            new_path = os.path.join(processed_dir, os.path.basename(file_path))
                            shutil_module.move(file_path, new_path)
                            logger.info(f"Moved processed file to: {new_path}")
                        except Exception as move_error:
                            logger.warning(f"Could not move file to processed folder: {move_error}")
                        
                        yield f"data: {json.dumps({'type': 'progress', 'message': f'Completed {file_name}', 'progress': int(((i + 1) / total_files) * 80), 'current_file': file_name, 'total_files': total_files, 'processed_files': processed_files})}\n\n"
                    else:
                        logger.warning(f"No content extracted from file: {file_path}")
                        yield f"data: {json.dumps({'type': 'progress', 'message': f'No content found in {file_name}', 'progress': int(((i + 1) / total_files) * 80), 'current_file': file_name, 'total_files': total_files, 'processed_files': processed_files})}\n\n"
                        document_ids.append(None)
                        
                except Exception as e:
                    logger.error(f"Error processing file {file_path}: {str(e)}")
                    yield f"data: {json.dumps({'type': 'error', 'message': f'Error processing {file_name}: {str(e)}', 'progress': int(((i + 1) / total_files) * 80), 'current_file': file_name, 'total_files': total_files, 'processed_files': processed_files})}\n\n"
                    document_ids.append(None)
                
                # Small delay to ensure UI updates
                await asyncio.sleep(0.1)
            
            # Final completion
            successful_count = len([doc_id for doc_id in document_ids if doc_id is not None])
            yield f"data: {json.dumps({'type': 'complete', 'message': f'Processing complete! Successfully processed {successful_count} out of {total_files} documents', 'progress': 100, 'total_files': total_files, 'processed_files': successful_count, 'document_ids': document_ids})}\n\n"
            
        except Exception as e:
            logger.error(f"Error in document processing stream: {str(e)}")
            yield f"data: {json.dumps({'type': 'error', 'message': f'Processing failed: {str(e)}', 'progress': 0})}\n\n"
    
    return StreamingResponse(
        generate_progress(),
        media_type="text/plain",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "Access-Control-Allow-Origin": "*",
            "Access-Control-Allow-Headers": "*",
        }
    )

@app.post("/process", response_model=ProcessResponse)
@monitor_performance
async def process_documents_endpoint(request: ProcessDocumentsRequest):
    """Process uploaded documents and create vector stores (legacy endpoint)"""
    try:
        document_ids = []
        all_documents = []
        
        # Load all documents
        for file_path in request.file_paths:
            if not os.path.exists(file_path):
                print(f"Warning: File not found: {file_path}")
                document_ids.append(None)  # Add None for missing files
                continue
            
            try:
                docs = load_document(file_path)
                if docs:
                    all_documents.extend(docs)
                    
                    # Generate unique ID
                    doc_id = str(uuid4())
                    
                    # Process this individual document into main ChromaDB collection
                    chunk_count = process_documents(docs, doc_id)
                    
                    # Store metadata
                    document_metadata[doc_id] = {
                        "filename": os.path.basename(file_path),
                        "processed_at": time.time(),
                        "chunk_count": chunk_count,
                        "file_size": os.path.getsize(file_path)
                    }
                    document_ids.append(doc_id)
                    
                    # Move processed file to processed-documents folder
                    try:
                        # Create processed-documents directory if it doesn't exist
                        uploads_dir = os.path.dirname(file_path)
                        processed_dir = os.path.join(uploads_dir, "processed-documents")
                        os.makedirs(processed_dir, exist_ok=True)
                        
                        # Move the file
                        new_path = os.path.join(processed_dir, os.path.basename(file_path))
                        shutil_module.move(file_path, new_path)
                        logger.info(f"Moved processed file to: {new_path}")
                    except Exception as move_error:
                        logger.warning(f"Could not move file to processed folder: {move_error}")
                else:
                    print(f"Warning: No content extracted from file: {file_path}")
                    document_ids.append(None)
            except Exception as e:
                print(f"Error processing file {file_path}: {str(e)}")
                document_ids.append(None)
        
        # Count successful documents (non-None IDs)
        successful_count = len([doc_id for doc_id in document_ids if doc_id is not None])
        
        return ProcessResponse(
            document_ids=document_ids,
            message=f"Successfully processed {successful_count} out of {len(request.file_paths)} documents"
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing documents: {str(e)}")

@app.post("/chat", response_model=ChatResponse)
@monitor_performance
async def chat_endpoint(request: ChatRequest):
    """Chat with the RAG system using production-grade retrieval"""
    logger.info(f"Chat request received for {len(request.documents)} documents")
    
    try:
        if not request.documents:
            logger.warning("Chat request received with no documents")
            raise HTTPException(status_code=400, detail="No documents provided")
        
        # Collect relevant documents from ALL uploaded documents
        all_relevant_docs = []
        valid_doc_ids = []
        
        # Create conversation-aware query for better retrieval
        conversation_context = ""
        if request.conversationHistory:
            # Get last few exchanges for context
            recent_history = request.conversationHistory[-4:]  # Last 2 Q&A pairs
            conversation_context = "\n".join([f"{msg['role']}: {msg['content']}" for msg in recent_history])
        
        # Enhanced query that includes conversation context
        enhanced_query = request.message
        if conversation_context:
            enhanced_query = f"Previous conversation:\n{conversation_context}\n\nCurrent question: {request.message}"
        
        # Get main vector store
        vectorstore = get_main_vectorstore()
        
        # Build metadata filter for requested documents
        valid_doc_ids = [doc_id for doc_id in request.documents if doc_id and doc_id in document_metadata]
        
        if not valid_doc_ids:
            logger.error(f"No valid documents found from IDs: {request.documents}")
            raise HTTPException(status_code=404, detail="No valid documents found")
        
        # Use ChromaDB's similarity search with metadata filtering
        # First, retrieve more candidates than needed for reranking
        filter_dict = {"document_id": {"$in": valid_doc_ids}}
        
        try:
            # Retrieve more candidates for reranking based on multiplier
            num_candidates = ChromaDBConfig.SEARCH_KWARGS["k"] * ChromaDBConfig.RERANK_MULTIPLIER
            candidate_docs = vectorstore.similarity_search(
                enhanced_query,
                k=num_candidates,
                filter=filter_dict
            )
            
            logger.info(f"Retrieved {len(candidate_docs)} candidate chunks for reranking")
            
            # Rerank using cross-encoder if we have candidates
            if candidate_docs:
                # Prepare pairs for cross-encoder
                pairs = [[request.message, doc.page_content] for doc in candidate_docs]
                
                # Get cross-encoder scores
                scores = cross_encoder.predict(pairs)
                
                # Sort documents by cross-encoder scores
                doc_score_pairs = list(zip(candidate_docs, scores))
                doc_score_pairs.sort(key=lambda x: x[1], reverse=True)
                
                # Filter by minimum score and take top k
                filtered_docs = [(doc, score) for doc, score in doc_score_pairs 
                               if score >= ChromaDBConfig.MIN_RERANK_SCORE]
                
                # If no docs meet minimum score, take top k anyway
                if not filtered_docs:
                    logger.warning(f"No chunks scored above {ChromaDBConfig.MIN_RERANK_SCORE}, using top {ChromaDBConfig.RERANK_TOP_K} chunks")
                    relevant_docs = [doc for doc, score in doc_score_pairs[:ChromaDBConfig.RERANK_TOP_K]]
                else:
                    # Take top k documents after filtering
                    relevant_docs = [doc for doc, score in filtered_docs[:ChromaDBConfig.RERANK_TOP_K]]
                
                logger.info(f"Reranked to {len(relevant_docs)} most relevant chunks")
                logger.info(f"Top reranking scores: {[score for _, score in doc_score_pairs[:5]]}")
                
                # Log if we filtered out low-scoring documents
                if filtered_docs and len(filtered_docs) < len(doc_score_pairs):
                    logger.info(f"Filtered out {len(doc_score_pairs) - len(filtered_docs)} chunks with scores below {ChromaDBConfig.MIN_RERANK_SCORE}")
            else:
                relevant_docs = candidate_docs
            
        except Exception as e:
            logger.error(f"Error during similarity search or reranking: {str(e)}")
            raise HTTPException(status_code=500, detail="Error retrieving documents")
        
        # Check if we found any relevant documents
        if not relevant_docs:
            return ChatResponse(
                response="I don't have any documents to search through. Please upload and process some documents first.",
                sources=[]
            )
        
        # Production-grade context creation with source attribution
        context_parts = []
        for i, doc in enumerate(relevant_docs):
            # Add source information for better traceability
            source = doc.metadata.get("source", "Unknown")
            source_name = source.split("/")[-1] if "/" in source else source.split("\\")[-1] if "\\" in source else source
            
            context_parts.append(f"[Source: {source_name}]\n{doc.page_content}")
        
        context = "\n\n---\n\n".join(context_parts)
        
        # Build conversation-aware prompt
        conversation_history_text = ""
        if request.conversationHistory:
            recent_history = request.conversationHistory[-4:]  # Last 2 Q&A pairs
            conversation_history_text = "\n".join([f"{msg['role'].title()}: {msg['content']}" for msg in recent_history])
        
        # Production-grade prompt with conversation awareness
        system_prompt = """You are a professional AI assistant that provides accurate, concise answers based on document context.

CRITICAL INSTRUCTIONS:
1. Answer ONLY using information from the provided context
2. Use conversation history to understand references like "last method", "that approach", etc.
3. Never mention sources or document names in your response
4. If multiple sources contain relevant info, synthesize them coherently
5. If context is insufficient or irrelevant, simply state: "I don't have information about that in the provided documents."
6. Be precise and direct - avoid unnecessary explanations about what you don't know
7. Do NOT add external knowledge or assumptions
8. Keep responses concise and to the point

Provide a direct answer using ONLY the document context provided. Use the conversation history to understand contextual references."""
        
        user_message = f"""Context from uploaded documents:

{context}

---

{"Conversation History:\n" + conversation_history_text + "\n\n---\n\n" if conversation_history_text else ""}Current Question: {request.message}"""
        
        prompt = ChatPromptTemplate.from_messages([
            ("system", system_prompt),
            ("user", user_message)
        ])
        
        # Generate response without timeout restrictions
        try:
            chain = prompt | llm | StrOutputParser()
            response = await asyncio.to_thread(chain.invoke, {"context": context, "question": request.message})
            logger.info(f"Generated response of length {len(response)}")
        except Exception as e:
            logger.error(f"LLM response generation failed: {str(e)}")
            raise HTTPException(status_code=500, detail="Response generation failed")
        
        # Get source information
        sources = [doc.metadata.get("source", "Unknown") for doc in relevant_docs]
        
        # Clean and deduplicate sources
        unique_sources = list(set(sources))
        logger.info(f"Response generated using {len(unique_sources)} sources")
        
        return ChatResponse(
            response=response,
            sources=unique_sources
        )
        
    except HTTPException:
        raise  # Re-raise HTTP exceptions as-is
    except Exception as e:
        logger.error(f"Unexpected error in chat endpoint: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail="An unexpected error occurred while generating the response")

@app.get("/documents")
async def list_documents():
    """List all processed documents with metadata"""
    try:
        documents_info = []
        
        for doc_id, metadata in document_metadata.items():
            doc_info = {
                "id": doc_id,
                **metadata
            }
            documents_info.append(doc_info)
        
        return {
            "documents": documents_info,
            "total_count": len(documents_info),
            "total_chunks": sum(doc.get("chunk_count", 0) for doc in documents_info),
            "total_size_bytes": sum(doc.get("file_size", 0) for doc in documents_info)
        }
    except Exception as e:
        logger.error(f"Error listing documents: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error listing documents: {str(e)}")

@app.delete("/documents/{document_id}")
async def delete_document(document_id: str):
    """Delete a processed document set"""
    try:
        if document_id in document_metadata:
            # Get the main vector store
            vectorstore = get_main_vectorstore()
            
            # Delete all chunks with this document_id from ChromaDB
            try:
                # Get the collection directly from ChromaDB client
                collection = chroma_client.get_collection(name=ChromaDBConfig.MAIN_COLLECTION)
                
                # Delete documents with matching document_id
                collection.delete(
                    where={"document_id": document_id}
                )
                
                logger.info(f"Deleted all chunks for document {document_id} from main collection")
            except Exception as e:
                logger.error(f"Error deleting chunks from ChromaDB: {str(e)}")
            
            # Remove from metadata
            del document_metadata[document_id]
            
            return {"message": "Document deleted successfully"}
        else:
            raise HTTPException(status_code=404, detail="Document not found")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error deleting document: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
